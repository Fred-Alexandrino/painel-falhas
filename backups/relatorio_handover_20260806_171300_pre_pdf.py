# -*- coding: utf-8 -*-
"""
relatorio_handover.py
─────────────────────────────────────────────────────────────────────────
Geração do "Relatório de Handover" (.docx) por OS, no padrão visual
Grid Co. (navy #191528 / lime #A9DB21), a partir de uma atividade do
Painel de Atividades (dict com as chaves de ATIV_HEADERS_JSON do app.py).

Todos os campos objetivos (cliente, usina, equipamento, datas, status,
histórico) vêm direto dos dados da OS — determinístico, nunca inventado.
O único trecho opcional gerado por IA é o resumo executivo (texto corrido
formal), passado já pronto via `resumo_ia`; se vier vazio, a seção é
simplesmente omitida e o relatório sai só com os dados objetivos.
"""
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PALETA = {"navy": "191528", "secundaria": "504C63", "lime": "A9DB21",
          "verde": "2E7D32", "ambar": "B9770E", "vermelho": "B3261E", "cinza": "6B6B78"}

STATUS_LABEL = {
    "em processo": "Em Execução",
    "em revisão": "Em Verificação",
    "em revisao": "Em Verificação",
    "finalizada": "Concluída",
    "cancelada": "Cancelada",
}


def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexc)
    tcPr.append(sh)


def _borders(tbl, color="D9D9E0"):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for e in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement('w:' + e)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        b.append(el)
    tblPr.append(b)


def _novo_doc(titulo_txt, subtitulo_txt=""):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(1.9)
        s.right_margin = Cm(1.9)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    p = doc.add_paragraph()
    r = p.add_run(titulo_txt)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(PALETA["navy"])

    if subtitulo_txt:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(subtitulo_txt)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor.from_string(PALETA["cinza"])

    bar = doc.add_paragraph()
    pPr = bar._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single')
    bt.set(qn('w:sz'), '18')
    bt.set(qn('w:color'), PALETA["lime"])
    pbdr.append(bt)
    pPr.append(pbdr)
    return doc


def _secao(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), PALETA["navy"])
    pPr.append(sh)
    r = p.add_run("  " + txt + "  ")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("FFFFFF")
    return p


def _paragrafo(doc, texto, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texto or "—")
    r.italic = italic
    return p


def _tabela_navy(doc, cabecalhos, linhas, larguras_cm=None):
    tb = doc.add_table(rows=1, cols=len(cabecalhos))
    _borders(tb)
    for i, h in enumerate(tb.rows[0].cells):
        h.text = ""
        r = h.paragraphs[0].add_run(cabecalhos[i])
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        r.font.size = Pt(9.5)
        _shade(h, PALETA["navy"])
    for linha in linhas:
        cs = tb.add_row().cells
        for j, v in enumerate(linha):
            cs[j].text = ""
            r = cs[j].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
    if larguras_cm:
        for i, w in enumerate(larguras_cm):
            tb.columns[i].width = Cm(w)
    return tb


def _limpar_historico(historico_raw, max_linhas=15):
    """Quebra o histórico bruto (texto livre, uma entrada por linha) em
    linhas legíveis, sem alterar o conteúdo original de cada entrada."""
    if not historico_raw:
        return []
    linhas = [l.strip(" -•\t") for l in re.split(r"[\r\n]+", historico_raw) if l.strip()]
    return linhas[:max_linhas]


def gerar_handover_docx(atividade, resumo_ia=""):
    """
    Monta o Relatório de Handover (.docx) no padrão visual Grid Co. a
    partir de uma atividade (dict, chaves = ATIV_HEADERS_JSON do app.py).

    `resumo_ia`: texto do resumo executivo já gerado (opcional). Se vazio,
    a seção "Resumo Executivo" é omitida — o relatório nunca falha por
    causa da IA.

    Retorna um BytesIO posicionado no início, pronto para send_file.
    """
    cliente = (atividade.get("cliente") or "").strip()
    usina = (atividade.get("usina") or "").strip()
    numero_os = (atividade.get("numeroOS") or "").strip()
    equipamento = (atividade.get("equipamento") or "").strip()
    descricao = (atividade.get("descricao") or "").strip()
    responsavel = (atividade.get("responsavel") or "").strip()
    status_raw = (atividade.get("statusOS") or atividade.get("status") or "").strip()
    status_label = STATUS_LABEL.get(status_raw.lower(), status_raw or "—")
    data_criacao = (atividade.get("dataCriacao") or "").strip()
    data_conclusao = (atividade.get("dataConclusao") or "").strip()
    observacoes = (atividade.get("observacoesOS") or "").strip()
    historico_raw = atividade.get("historico") or ""

    doc = _novo_doc("Relatório de Handover", f"{cliente} — {usina}" if cliente or usina else "")

    _secao(doc, "1. Dados Gerais")
    _tabela_navy(doc, ["Campo", "Informação"], [
        ["Cliente", cliente or "—"],
        ["Usina", usina or "—"],
        ["Equipamento", equipamento or "—"],
        ["Nº da OS", numero_os or "—"],
        ["Responsável (Grid Co.)", responsavel or "—"],
        ["Status", status_label],
        ["Data de abertura", data_criacao or "—"],
        ["Data de conclusão", data_conclusao or "—"],
    ], larguras_cm=[5.5, 10.5])

    _secao(doc, "2. Descrição do Serviço")
    _paragrafo(doc, descricao or "Sem descrição registrada.")

    proxima_secao = 3
    if resumo_ia:
        _secao(doc, f"{proxima_secao}. Resumo Executivo")
        for paragrafo in resumo_ia.split("\n\n"):
            if paragrafo.strip():
                _paragrafo(doc, paragrafo.strip())
        proxima_secao += 1

    _secao(doc, f"{proxima_secao}. Histórico de Execução")
    linhas_hist = _limpar_historico(historico_raw)
    if linhas_hist:
        for linha in linhas_hist:
            _paragrafo(doc, "• " + linha)
    else:
        _paragrafo(doc, "Sem histórico detalhado registrado.", italic=True)
    proxima_secao += 1

    if observacoes:
        _secao(doc, f"{proxima_secao}. Observações")
        _paragrafo(doc, observacoes)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Grid Co. — Serviços Especializados de O&M")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(PALETA["cinza"])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
