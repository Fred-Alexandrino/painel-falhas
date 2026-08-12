# -*- coding: utf-8 -*-
"""
relatorio_handover_usina_docx.py
─────────────────────────────────────────────────────────────────────────
Gera o "Relatório de Handover — Operação de Ativos" (.docx) a partir do
MODELO REAL da Grid Co. (templates/handover_grid_template.docx — enviado
pelo Fred, com a capa, cabeçalho, estilos e logo verdadeiros), substituindo
só o conteúdo variável (usina, cliente, datas, equipe, textos das etapas,
lista de ativos do handover, capacitação, conclusão, quadro de revisões)
e anexando a Punch List como uma seção final em paisagem.

Diferente de relatorio_handover_usina.py (que RECRIA o documento do zero
via reportlab, uma aproximação visual da identidade Grid Co.), este
módulo edita o arquivo .docx de verdade — mesma capa (foto + logo),
mesmo cabeçalho, mesmas fontes/estilos, porque é literalmente o mesmo
arquivo, só com o conteúdo trocado.

Limitação conhecida: o Word não tem como "colar" as páginas de um PDF
externo dentro de si (diferente de PDF-em-PDF). A seção "3.4 Ordens de
Serviço - Handover" fica só com o título — o PDF exportado da Fracttal
continua sendo um anexo separado, não embutido neste .docx.
"""
import copy
import os
import subprocess
import tempfile
from io import BytesIO

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import pdfplumber
from pypdf import PdfReader, PdfWriter

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "handover_grid_template.docx")

VERDE_PUNCHLIST = "2E7D32"
BORDA_PUNCHLIST = "D9D9E0"
ZEBRA_PUNCHLIST = "F0F1F4"


# ── Helpers de edição de parágrafos existentes ──────────────────────────

def _definir_texto_paragrafo(paragrafo, novo_texto):
    """Substitui o texto visível de um parágrafo por `novo_texto`,
    mantendo a formatação do primeiro run (fonte, negrito, cor etc.) —
    funciona mesmo que o texto original esteja fragmentado em vários
    runs (comum em docx editados no Word)."""
    if paragrafo is None:
        return
    if not paragrafo.runs:
        paragrafo.add_run(novo_texto)
        return
    paragrafo.runs[0].text = novo_texto
    for r in paragrafo.runs[1:]:
        r.text = ""


def _achar_paragrafo(doc, contendo, a_partir_de=0):
    for i, p in enumerate(doc.paragraphs):
        if i < a_partir_de:
            continue
        if contendo in p.text:
            return p
    return None


def _clonar_e_inserir_antes(paragrafo_modelo, novo_texto):
    """Clona a formatação de `paragrafo_modelo` num novo parágrafo
    inserido logo ANTES dele, com `novo_texto`. Retorna o novo parágrafo
    (um docx.text.paragraph.Paragraph de verdade)."""
    novo_p_elem = copy.deepcopy(paragrafo_modelo._p)
    paragrafo_modelo._p.addprevious(novo_p_elem)
    novo_paragrafo = Paragraph(novo_p_elem, paragrafo_modelo._parent)
    if novo_paragrafo.runs:
        novo_paragrafo.runs[0].text = novo_texto
        for r in novo_paragrafo.runs[1:]:
            r.text = ""
    else:
        novo_paragrafo.add_run(novo_texto)
    return novo_paragrafo


def _substituir_bloco_paragrafos(doc, ancora_primeiro, ancora_ultimo, novos_textos):
    """
    Localiza o intervalo de parágrafos entre o que contém `ancora_primeiro`
    e o que contém `ancora_ultimo` (nessa ordem, no mesmo bloco), e troca
    esse intervalo inteiro pelos `novos_textos` (uma linha por item),
    reaproveitando a formatação do primeiro parágrafo do bloco original.
    Usado pros blocos de tamanho variável: Equipe, lista "Handover – X" e
    Capacitação. Se não achar as âncoras, não faz nada (evita quebrar o
    documento se o modelo mudar)."""
    paragrafos = doc.paragraphs
    idx_inicio = idx_fim = None
    for i, p in enumerate(paragrafos):
        if idx_inicio is None and ancora_primeiro in p.text:
            idx_inicio = i
        if idx_inicio is not None and i >= idx_inicio and ancora_ultimo in p.text:
            idx_fim = i
            break
    if idx_inicio is None or idx_fim is None:
        return

    modelo = paragrafos[idx_inicio]
    elementos_originais = [paragrafos[i]._p for i in range(idx_inicio, idx_fim + 1)]

    textos = [t for t in novos_textos if t and t.strip()] or [""]
    for texto in textos:
        _clonar_e_inserir_antes(modelo, texto.strip())

    for el in elementos_originais:
        el.getparent().remove(el)


# ── Punch List (seção final em paisagem) ────────────────────────────────

def _sombrear_celula(celula, hex_cor):
    tcPr = celula._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_cor)
    tcPr.append(sh)


def _bordas_tabela(tabela, cor="D9D9E0"):
    """Aplica bordas finas em toda a grade da tabela — sem isso, o Word
    não desenha NENHUMA linha entre as células (fica sem cara de
    tabela nenhuma, só texto meio alinhado)."""
    tblPr = tabela._tbl.tblPr
    bordas = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), cor)
        bordas.append(el)
    tblPr.append(bordas)


def _definir_largura_coluna(tabela, indice, largura_cm):
    """No python-docx, a largura de coluna só é respeitada de verdade
    no Word se for setada em CADA célula da coluna, não só no objeto
    Table — por isso o define_col_widths do table sozinho não bastava
    e as colunas saíam com larguras arbitrárias/erradas."""
    for linha in tabela.rows:
        linha.cells[indice].width = Cm(largura_cm)


def _adicionar_secao_punchlist(doc, punch_list):
    """Adiciona uma nova seção em paisagem no final do documento com a
    tabela de Punch List (mesmo padrão visual: cabeçalho verde, zebra,
    bordas finas, colunas com largura fixa proporcional ao conteúdo)."""
    nova_secao = doc.add_section(WD_SECTION.NEW_PAGE)
    nova_secao.orientation = 1  # WD_ORIENT.LANDSCAPE
    largura_antiga, altura_antiga = nova_secao.page_width, nova_secao.page_height
    nova_secao.page_width, nova_secao.page_height = altura_antiga, largura_antiga
    nova_secao.left_margin = nova_secao.right_margin = Cm(1.2)
    nova_secao.top_margin = nova_secao.bottom_margin = Cm(1.6)
    # a seção nova herda o mesmo cabeçalho — como é paisagem e a punch
    # list é a última página, "linkar ao anterior" (padrão) já é o
    # comportamento correto pro modelo Grid Co. (o cabeçalho não aparece
    # nessa página no modelo de referência, mas manter é inofensivo)

    titulo = doc.add_paragraph()
    r = titulo.add_run("Punch List")
    r.bold = True
    r.font.size = Pt(15)
    titulo.paragraph_format.space_after = Pt(10)

    cabecalho = ["CLIENTE", "USINA", "CLUSTER", "ATIVO", "CRITICIDADE",
                 "STATUS", "ANORMALIDADE", "RECOMENDAÇÕES", "RESPONSÁVEL"]
    # larguras em cm, somando ~25.5cm (A4 paisagem ~29.7cm - 2x1.2cm margem)
    larguras_cm = [2.6, 2.8, 2.2, 2.6, 2.2, 2.2, 4.8, 4.8, 2.7]

    tabela = doc.add_table(rows=1, cols=len(cabecalho))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    _bordas_tabela(tabela)

    for i, texto in enumerate(cabecalho):
        celula = tabela.rows[0].cells[i]
        celula.text = ""
        p = celula.paragraphs[0]
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        _sombrear_celula(celula, VERDE_PUNCHLIST)

    for idx, item in enumerate(punch_list):
        linha = tabela.add_row()
        valores = [
            item.get("cliente", ""), item.get("usina", ""), item.get("cluster", ""),
            item.get("ativo", ""), item.get("criticidade", ""), item.get("status", "PENDENTE"),
            item.get("anormalidade", ""), item.get("recomendacoes", ""), item.get("responsavel", ""),
        ]
        for i, valor in enumerate(valores):
            celula = linha.cells[i]
            celula.text = ""
            p = celula.paragraphs[0]
            r = p.add_run(valor or "—")
            r.font.size = Pt(8.5)
            if idx % 2 == 1:
                _sombrear_celula(celula, ZEBRA_PUNCHLIST)

    for i, largura in enumerate(larguras_cm):
        _definir_largura_coluna(tabela, i, largura)


# ── Função principal ─────────────────────────────────────────────────

def gerar_handover_usina_docx(dados):
    """
    Gera o Relatório de Handover (.docx) a partir do modelo real da
    Grid Co., substituindo o conteúdo variável de `dados` (mesmas
    chaves usadas em relatorio_handover_usina.py: cliente, usina,
    localizacao, dataInicio, dataFim, equipe, planejamento,
    documentacaoEntregue, equipamentosHandover, capacitacao, conclusao,
    revisao, punchList).

    Retorna um BytesIO posicionado no início, pronto para send_file.
    """
    doc = Document(TEMPLATE_PATH)

    cliente = dados.get("cliente", "")
    usina = dados.get("usina", "")
    equipamentos = dados.get("equipamentosHandover", [])
    equipe = dados.get("equipe", [])
    capacitacao = dados.get("capacitacao", [])
    revisao = dados.get("revisao", {})

    # Capa
    p_capa = _achar_paragrafo(doc, "UFV GUAJIRU")
    _definir_texto_paragrafo(p_capa, f"UFV {usina.upper()} – {cliente.upper()}")

    # 1. Objetivo
    p_obj = _achar_paragrafo(doc, "Este relatório tem como objetivo")
    _definir_texto_paragrafo(p_obj, (
        f"Este relatório tem como objetivo documentar o processo de handover da usina "
        f"fotovoltaica {cliente} – {usina} da equipe de Construção/EPC para a equipe de "
        f"Operação e Manutenção (O&M). O handover foi realizado com o objetivo de garantir "
        f"a transferência de responsabilidades de maneira eficiente, segura e devidamente "
        f"documentada, assegurando que a usina esteja em condições plenas de operação."
    ))

    # 2. Informações Gerais
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Nome da Usina:"), f"Nome da Usina: {usina}")
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Localização:"),
                              f"Localização: {dados.get('localizacao', '') or '—'}")
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Data de Início do Handover:"),
                              f"Data de Início do Handover: {dados.get('dataInicio', '') or '—'}")
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Data de Conclusão do Handover:"),
                              f"Data de Conclusão do Handover: {dados.get('dataFim', '') or '—'}")

    # Equipe de O&M (bloco de tamanho variável)
    _substituir_bloco_paragrafos(doc, "Supervisor: Fred Alexandrino", "Mantenedor Isake Costa",
                                  equipe or ["—"])

    # 3.1 Planejamento
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Foi elaborado um cronograma"),
                              dados.get("planejamento") or "—")

    # 3.2 Documentação Entregue
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Foram entregues à equipe de O&M"),
                              dados.get("documentacaoEntregue") or "—")

    # 3.3 Inspeções e Testes (intro genérica com o nome da usina) + lista
    p_insp = _achar_paragrafo(doc, "Foram realizadas inspeções físicas")
    _definir_texto_paragrafo(p_insp, (
        f"Foram realizadas inspeções físicas, funcionais e termográficas em todos os "
        f"sistemas e componentes principais da UFV {usina}, com o objetivo de verificar a "
        f"integridade, o desempenho operacional e eventuais falhas que possam comprometer "
        f"a segurança ou a eficiência da planta. Os resultados são apresentados nos tópicos "
        f"a seguir."
    ))
    _substituir_bloco_paragrafos(
        doc, "Handover – Cabine de Medição", "Handover – Transformador de Potência",
        [f"Handover – {eq}" for eq in equipamentos] or ["—"])

    # 4. Capacitação da Equipe (bloco de tamanho variável)
    _substituir_bloco_paragrafos(doc, "Procedimentos de segurança.", "Manutenção preventiva e corretiva.",
                                  capacitacao or ["—"])

    # 5. Conclusão
    _definir_texto_paragrafo(_achar_paragrafo(doc, "O handover da UFV"), dados.get("conclusao") or "—")

    # 6. Quadro de Revisões (tabela já existe no modelo — só preenche a 2ª linha)
    if doc.tables:
        tabela_revisoes = doc.tables[0]
        if len(tabela_revisoes.rows) >= 2:
            valores = [
                revisao.get("revisao", "00"), revisao.get("edicao", "Emissão inicial"),
                revisao.get("elaborador", "—"), revisao.get("verificador", "—"),
                revisao.get("aprovador", "—"), revisao.get("data", "—"),
            ]
            for i, valor in enumerate(valores):
                if i < len(tabela_revisoes.rows[1].cells):
                    celula = tabela_revisoes.rows[1].cells[i]
                    for p in celula.paragraphs:
                        for r in p.runs:
                            r.text = ""
                    if celula.paragraphs[0].runs:
                        celula.paragraphs[0].runs[0].text = valor
                    else:
                        celula.paragraphs[0].add_run(valor)

    # Quebra de página logo após o título "3.4 Ordens de Serviço - Handover"
    # — garante que a seção 4 comece numa página nova, o que deixa o ponto
    # de corte (pra inserir o PDF da Fracttal, ver gerar_handover_usina_completo)
    # sem ambiguidade: tudo até a última página que contém esse título vira
    # a "parte 1", o resto vira "parte 2".
    p_ordens = _achar_paragrafo(doc, "Ordens de Serviço - Handover")
    if p_ordens is not None:
        novo_p_elem = OxmlElement("w:p")
        p_ordens._p.addnext(novo_p_elem)
        paragrafo_quebra = Paragraph(novo_p_elem, p_ordens._parent)
        paragrafo_quebra.add_run().add_break(WD_BREAK.PAGE)

    # Punch List (seção nova em paisagem, no final)
    _adicionar_secao_punchlist(doc, dados.get("punchList", []))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Conversão pra PDF + merge com o PDF da Fracttal ─────────────────────

def _converter_docx_para_pdf(docx_bytes):
    """Converte bytes de um .docx pra PDF via LibreOffice headless. Cada
    chamada usa um perfil de usuário temporário isolado, pra não haver
    conflito entre requisições concorrentes na mesma VM."""
    with tempfile.TemporaryDirectory(prefix="handover_lo_") as tmpdir:
        docx_path = os.path.join(tmpdir, "entrada.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        perfil = os.path.join(tmpdir, "perfil_lo")
        os.makedirs(perfil, exist_ok=True)

        resultado = subprocess.run(
            ["soffice", "--headless", "--norestore",
             f"-env:UserInstallation=file://{perfil}",
             "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
            capture_output=True, timeout=90,
        )
        pdf_path = os.path.join(tmpdir, "entrada.pdf")
        if not os.path.exists(pdf_path):
            erro = (resultado.stderr or b"").decode(errors="ignore")[:800]
            raise RuntimeError(f"Falha ao converter o .docx para PDF (LibreOffice): {erro}")
        with open(pdf_path, "rb") as f:
            return f.read()


def _dividir_pdf_no_marcador(pdf_bytes, texto_marcador="Ordens de Serviço - Handover"):
    """Divide o PDF em duas partes: tudo até a ÚLTIMA página que contém
    `texto_marcador` (inclusive) vira a parte 1, o resto vira a parte 2.
    Como gerar_handover_usina_docx já insere uma quebra de página logo
    depois desse título, essa página nunca tem conteúdo da seção 4 junto
    — o corte é limpo. Se o marcador não for encontrado (modelo mudou),
    não divide: tudo vira parte 1 e a parte 2 fica vazia."""
    reader = PdfReader(BytesIO(pdf_bytes))
    pagina_split = None
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            if texto_marcador in (page.extract_text() or ""):
                pagina_split = i

    parte1, parte2 = PdfWriter(), PdfWriter()
    for i, page in enumerate(reader.pages):
        alvo = parte1 if (pagina_split is None or i <= pagina_split) else parte2
        alvo.add_page(page)

    buf1, buf2 = BytesIO(), BytesIO()
    parte1.write(buf1)
    parte2.write(buf2)
    buf1.seek(0)
    buf2.seek(0)
    return buf1.getvalue(), buf2.getvalue()


def gerar_handover_usina_completo(dados, fracttal_pdf_bytes=None):
    """
    Gera o Relatório de Handover completo.

    Sem PDF da Fracttal: retorna (bytes_docx, "docx") — o documento
    editável, direto do modelo real.

    Com PDF da Fracttal: converte o .docx pra PDF (LibreOffice), corta
    logo após o título "3.4 Ordens de Serviço - Handover", insere as
    páginas do PDF da Fracttal ali no meio, e devolve
    (bytes_pdf_final, "pdf") — documento + OS mesclados, prontos pra
    enviar ao cliente.
    """
    docx_bytes = gerar_handover_usina_docx(dados).read()
    if not fracttal_pdf_bytes:
        return docx_bytes, "docx"

    pdf_bytes = _converter_docx_para_pdf(docx_bytes)
    parte1_bytes, parte2_bytes = _dividir_pdf_no_marcador(pdf_bytes)

    writer = PdfWriter()
    for chunk in (parte1_bytes, fracttal_pdf_bytes, parte2_bytes):
        if not chunk:
            continue
        for page in PdfReader(BytesIO(chunk)).pages:
            writer.add_page(page)

    saida = BytesIO()
    writer.write(saida)
    saida.seek(0)
    return saida.read(), "pdf"
