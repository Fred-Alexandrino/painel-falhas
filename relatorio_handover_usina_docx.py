# -*- coding: utf-8 -*-
"""
relatorio_handover_usina_docx.py
─────────────────────────────────────────────────────────────────────────
Gera o "Relatório de Handover — Operação de Ativos" (.docx) a partir do
MODELO REAL da Grid Co. (templates/handover_grid_template.docx — enviado
pelo Fred, com a capa, cabeçalho, estilos e logo verdadeiros), substituindo
só o conteúdo variável (usina, cliente, datas, equipe, textos das etapas,
lista de ativos do handover, capacitação, conclusão, quadro de revisões)
e anexando a Punch List como uma seção final em paisagem.

Diferente de relatorio_handover_usina.py (que RECRIA o documento do zero
via reportlab, uma aproximação visual da identidade Grid Co.), este
módulo edita o arquivo .docx de verdade — mesma capa (foto + logo),
mesmo cabeçalho, mesmas fontes/estilos, porque é literalmente o mesmo
arquivo, só com o conteúdo trocado.

Limitação conhecida: o Word não tem como "colar" as páginas de um PDF
externo dentro de si (diferente de PDF-em-PDF). A seção "3.4 Ordens de
Serviço - Handover" fica só com o título — o PDF exportado da Fracttal
continua sendo um anexo separado, não embutido neste .docx.
"""
import copy
import os
import subprocess
import tempfile
from io import BytesIO

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import pdfplumber
from pypdf import PdfReader, PdfWriter

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "handover_grid_template.docx")

VERDE_PUNCHLIST = "2E7D32"
BORDA_PUNCHLIST = "D9D9E0"
ZEBRA_PUNCHLIST = "F0F1F4"


# ── Helpers de edição de parágrafos existentes ──────────────────────────

def _definir_texto_paragrafo(paragrafo, novo_texto):
    """Substitui o texto visível de um parágrafo por `novo_texto`,
    mantendo a formatação do primeiro run (fonte, negrito, cor etc.) —
    funciona mesmo que o texto original esteja fragmentado em vários
    runs (comum em docx editados no Word)."""
    if paragrafo is None:
        return
    if not paragrafo.runs:
        paragrafo.add_run(novo_texto)
        return
    paragrafo.runs[0].text = novo_texto
    for r in paragrafo.runs[1:]:
        r.text = ""


def _achar_paragrafo(doc, contendo, a_partir_de=0):
    for i, p in enumerate(doc.paragraphs):
        if i < a_partir_de:
            continue
        if contendo in p.text:
            return p
    return None


def _clonar_e_inserir_antes(paragrafo_modelo, novo_texto):
    """Clona a formatação de `paragrafo_modelo` num novo parágrafo
    inserido logo ANTES dele, com `novo_texto`. Retorna o novo parágrafo
    (um docx.text.paragraph.Paragraph de verdade)."""
    novo_p_elem = copy.deepcopy(paragrafo_modelo._p)
    paragrafo_modelo._p.addprevious(novo_p_elem)
    novo_paragrafo = Paragraph(novo_p_elem, paragrafo_modelo._parent)
    if novo_paragrafo.runs:
        novo_paragrafo.runs[0].text = novo_texto
        for r in novo_paragrafo.runs[1:]:
            r.text = ""
    else:
        novo_paragrafo.add_run(novo_texto)
    return novo_paragrafo


def _substituir_bloco_paragrafos(doc, ancora_primeiro, ancora_ultimo, novos_textos):
    """
    Localiza o intervalo de parágrafos entre o que contém `ancora_primeiro`
    e o que contém `ancora_ultimo` (nessa ordem, no mesmo bloco), e troca
    esse intervalo inteiro pelos `novos_textos` (uma linha por item),
    reaproveitando a formatação do primeiro parágrafo do bloco original.
    Usado pros blocos de tamanho variável: Equipe, lista "Handover – X" e
    Capacitação. Se não achar as âncoras, não faz nada (evita quebrar o
    documento se o modelo mudar)."""
    paragrafos = doc.paragraphs
    idx_inicio = idx_fim = None
    for i, p in enumerate(paragrafos):
        if idx_inicio is None and ancora_primeiro in p.text:
            idx_inicio = i
        if idx_inicio is not None and i >= idx_inicio and ancora_ultimo in p.text:
            idx_fim = i
            break
    if idx_inicio is None or idx_fim is None:
        return

    modelo = paragrafos[idx_inicio]
    elementos_originais = [paragrafos[i]._p for i in range(idx_inicio, idx_fim + 1)]

    textos = [t for t in novos_textos if t and t.strip()] or [""]
    for texto in textos:
        _clonar_e_inserir_antes(modelo, texto.strip())

    for el in elementos_originais:
        el.getparent().remove(el)


# ── Punch List (seção final em paisagem) ────────────────────────────────

def _sombrear_celula(celula, hex_cor):
    tcPr = celula._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_cor)
    tcPr.append(sh)


def _bordas_tabela(tabela, cor="D9D9E0"):
    """Aplica bordas finas em toda a grade da tabela — sem isso, o Word
    não desenha NENHUMA linha entre as células (fica sem cara de
    tabela nenhuma, só texto meio alinhado)."""
    tblPr = tabela._tbl.tblPr
    bordas = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), cor)
        bordas.append(el)
    tblPr.append(bordas)


def _definir_largura_coluna(tabela, indice, largura_cm):
    """Define a largura de uma coluna 'de verdade': atualiza tanto o
    tblGrid (a definição de colunas da tabela, que é o que o Word usa
    de fato quando w:tblLayout="fixed") quanto a largura de cada célula
    dessa coluna. Setar só a célula (como a versão antiga fazia) não
    tem efeito nenhum em tabelas com tblGrid explícito — caso do
    Quadro de Revisões, que veio pronto no modelo."""
    grid = tabela._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        cols = grid.findall(qn('w:gridCol'))
        if indice < len(cols):
            cols[indice].set(qn('w:w'), str(int(largura_cm * 566.929)))  # cm -> dxa
    for linha in tabela.rows:
        if indice < len(linha.cells):
            linha.cells[indice].width = Cm(largura_cm)


def _repetir_linha_cabecalho(linha):
    """Marca a linha como cabeçalho repetido em cada página nova da
    tabela — sem isso, no modelo anterior o cabeçalho só aparecia na
    primeira página e a linha de dados que caía bem na quebra de página
    ficava cortada ao meio, sem contexto de coluna nenhum."""
    trPr = linha._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


def _impedir_quebra_de_linha(linha):
    """Impede que UMA linha da tabela seja dividida entre duas páginas
    — sem isso, uma linha de punch list mais longa (texto grande em
    Anormalidade/Recomendações) podia começar numa página e terminar
    na seguinte, cortando o texto no meio de forma ilegível."""
    trPr = linha._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def _adicionar_secao_punchlist(doc, punch_list):
    """Adiciona uma nova seção em paisagem no final do documento com a
    tabela de Punch List (mesmo padrão visual: cabeçalho verde, zebra,
    bordas finas, colunas com largura fixa proporcional ao conteúdo)."""
    nova_secao = doc.add_section(WD_SECTION.NEW_PAGE)
    nova_secao.orientation = 1  # WD_ORIENT.LANDSCAPE
    largura_antiga, altura_antiga = nova_secao.page_width, nova_secao.page_height
    nova_secao.page_width, nova_secao.page_height = altura_antiga, largura_antiga
    nova_secao.left_margin = nova_secao.right_margin = Cm(1.2)
    nova_secao.top_margin = nova_secao.bottom_margin = Cm(1.6)
    # a seção nova herda o mesmo cabeçalho — como é paisagem e a punch
    # list é a última página, "linkar ao anterior" (padrão) já é o
    # comportamento correto pro modelo Grid Co. (o cabeçalho não aparece
    # nessa página no modelo de referência, mas manter é inofensivo)

    titulo = doc.add_paragraph()
    r = titulo.add_run("Punch List")
    r.bold = True
    r.font.size = Pt(15)
    titulo.paragraph_format.space_after = Pt(10)
    # "Manter com o próximo" — sem isso, o LibreOffice às vezes empurra a
    # tabela inteira pra página seguinte (por causa do cabeçalho repetido +
    # "não quebrar linha"), deixando o título sozinho numa página em branco.
    titulo.paragraph_format.keep_with_next = True

    cabecalho = ["CLIENTE", "USINA", "CLUSTER", "ATIVO", "CRITICIDADE",
                 "STATUS", "ANORMALIDADE", "RECOMENDAÇÕES", "RESPONSÁVEL"]
    # larguras em cm, somando ~25.5cm (A4 paisagem ~29.7cm - 2x1.2cm margem)
    larguras_cm = [2.6, 2.8, 2.2, 2.6, 2.2, 2.2, 4.8, 4.8, 2.7]

    tabela = doc.add_table(rows=1, cols=len(cabecalho))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    _bordas_tabela(tabela)
    _repetir_linha_cabecalho(tabela.rows[0])
    _impedir_quebra_de_linha(tabela.rows[0])

    for i, texto in enumerate(cabecalho):
        celula = tabela.rows[0].cells[i]
        celula.text = ""
        p = celula.paragraphs[0]
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        _sombrear_celula(celula, VERDE_PUNCHLIST)

    for idx, item in enumerate(punch_list):
        linha = tabela.add_row()
        _impedir_quebra_de_linha(linha)
        valores = [
            item.get("cliente", ""), item.get("usina", ""), item.get("cluster", ""),
            item.get("ativo", ""), item.get("criticidade", ""), item.get("status", "PENDENTE"),
            item.get("anormalidade", ""), item.get("recomendacoes", ""), item.get("responsavel", ""),
        ]
        for i, valor in enumerate(valores):
            celula = linha.cells[i]
            celula.text = ""
            p = celula.paragraphs[0]
            r = p.add_run(valor or "—")
            r.font.size = Pt(8.5)
            if idx % 2 == 1:
                _sombrear_celula(celula, ZEBRA_PUNCHLIST)

    for i, largura in enumerate(larguras_cm):
        _definir_largura_coluna(tabela, i, largura)


# ── Função principal ─────────────────────────────────────────────────

def gerar_handover_usina_docx(dados, inserir_quebra_para_fracttal=False):
    """
    Gera o Relatório de Handover (.docx) a partir do modelo real da
    Grid Co., substituindo o conteúdo variável de `dados` (mesmas
    chaves usadas em relatorio_handover_usina.py: cliente, usina,
    localizacao, dataInicio, dataFim, equipe, planejamento,
    documentacaoEntregue, equipamentosHandover, capacitacao, conclusao,
    revisao, punchList).

    `inserir_quebra_para_fracttal`: só True quando o resultado vai ser
    convertido pra PDF e mesclado com o PDF da Fracttal em seguida (ver
    gerar_handover_usina_completo) — mantém o parágrafo em branco com
    quebra de página que já vem no modelo entre "3.4 Ordens de Serviço
    - Handover" e "4. CAPACITAÇÃO" (espaço reservado pra colar as
    páginas da Fracttal). Sem isso, esse parágrafo é removido, senão
    sobra uma página quase em branco no meio do .docx quando nenhum PDF
    é anexado.

    Retorna um BytesIO posicionado no início, pronto para send_file.
    """
    doc = Document(TEMPLATE_PATH)

    cliente = dados.get("cliente", "")
    usina = dados.get("usina", "")
    equipamentos = dados.get("equipamentosHandover", [])
    equipe = dados.get("equipe", [])
    capacitacao = dados.get("capacitacao", [])
    revisao = dados.get("revisao", {})

    # Capa
    p_capa = _achar_paragrafo(doc, "UFV GUAJIRU")
    _definir_texto_paragrafo(p_capa, f"UFV {usina.upper()} – {cliente.upper()}")

    # 1. Objetivo
    p_obj = _achar_paragrafo(doc, "Este relatório tem como objetivo")
    _definir_texto_paragrafo(p_obj, (
        f"Este relatório tem como objetivo documentar o processo de handover da usina "
        f"fotovoltaica {cliente} – {usina} da equipe de Construção/EPC para a equipe de "
        f"Operação e Manutenção (O&M). O handover foi realizado com o objetivo de garantir "
        f"a transferência de responsabilidades de maneira eficiente, segura e devidamente "
        f"documentada, assegurando que a usina esteja em condições plenas de operação."
    ))

    # 2. Informações Gerais
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Nome da Usina:"), f"Nome da Usina: {usina}")
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Localização:"),
                              f"Localização: {dados.get('localizacao', '') or '—'}")
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Data de Início do Handover:"),
                              f"Data de Início do Handover: {dados.get('dataInicio', '') or '—'}")
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Data de Conclusão do Handover:"),
                              f"Data de Conclusão do Handover: {dados.get('dataFim', '') or '—'}")

    # Equipe de O&M (bloco de tamanho variável)
    _substituir_bloco_paragrafos(doc, "Supervisor: Fred Alexandrino", "Mantenedor Isake Costa",
                                  equipe or ["—"])

    # 3.1 Planejamento
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Foi elaborado um cronograma"),
                              dados.get("planejamento") or "—")

    # 3.2 Documentação Entregue
    _definir_texto_paragrafo(_achar_paragrafo(doc, "Foram entregues à equipe de O&M"),
                              dados.get("documentacaoEntregue") or "—")

    # 3.3 Inspeções e Testes (intro genérica com o nome da usina) + lista
    p_insp = _achar_paragrafo(doc, "Foram realizadas inspeções físicas")
    _definir_texto_paragrafo(p_insp, (
        f"Foram realizadas inspeções físicas, funcionais e termográficas em todos os "
        f"sistemas e componentes principais da UFV {usina}, com o objetivo de verificar a "
        f"integridade, o desempenho operacional e eventuais falhas que possam comprometer "
        f"a segurança ou a eficiência da planta. Os resultados são apresentados nos tópicos "
        f"a seguir."
    ))
    _substituir_bloco_paragrafos(
        doc, "Handover – Cabine de Medição", "Handover – Transformador de Potência",
        [f"Handover – {eq}" for eq in equipamentos] or ["—"])

    # 4. Capacitação da Equipe (bloco de tamanho variável)
    _substituir_bloco_paragrafos(doc, "Procedimentos de segurança.", "Manutenção preventiva e corretiva.",
                                  capacitacao or ["—"])

    # 5. Conclusão
    _definir_texto_paragrafo(_achar_paragrafo(doc, "O handover da UFV"), dados.get("conclusao") or "—")

    # 6. Quadro de Revisões (tabela já existe no modelo — só preenche a 2ª linha)
    if doc.tables:
        tabela_revisoes = doc.tables[0]
        if len(tabela_revisoes.rows) >= 2:
            valores = [
                revisao.get("revisao", "00"), revisao.get("edicao", "Emissão inicial"),
                revisao.get("elaborador", "—"), revisao.get("verificador", "—"),
                revisao.get("aprovador", "—"), revisao.get("data", "—"),
            ]
            for i, valor in enumerate(valores):
                if i < len(tabela_revisoes.rows[1].cells):
                    celula = tabela_revisoes.rows[1].cells[i]
                    for p in celula.paragraphs:
                        for r in p.runs:
                            r.text = ""
                    if celula.paragraphs[0].runs:
                        celula.paragraphs[0].runs[0].text = valor
                    else:
                        celula.paragraphs[0].add_run(valor)
            # Fonte um pouco menor (cabeçalho e dados) — dá a folga que faltava
            # pra "Fred Alexandrino"/"Marcelo Martineli"/"13/08/2026" caberem
            # numa linha só sem precisar de colunas maiores que a largura útil
            # da página (16cm) comporta.
            for linha in tabela_revisoes.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9.5)
            # Larguras explícitas — no modelo original "Elaborador"/"Verificador"
            # quebravam a palavra ao meio (coluna estreita demais pro texto).
            tabela_revisoes.autofit = False
            larguras_revisoes = [1.7, 1.9, 3.3, 3.3, 3.4, 2.4]
            for i, largura in enumerate(larguras_revisoes):
                _definir_largura_coluna(tabela_revisoes, i, largura)
            tblW = tabela_revisoes._tbl.tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblW.set(qn('w:type'), 'dxa')
                tblW.set(qn('w:w'), str(int(sum(larguras_revisoes) * 566.929)))

    # Quebra de página logo após o título "3.4 Ordens de Serviço - Handover"
    # — garante que a seção 4 comece numa página nova, o que deixa o ponto
    # de corte (pra inserir o PDF da Fracttal, ver gerar_handover_usina_completo)
    # sem ambiguidade: tudo até a última página que contém esse título vira
    # a "parte 1", o resto vira "parte 2".
    p_ordens = _achar_paragrafo(doc, "Ordens de Serviço - Handover")
    if p_ordens is not None:
        novo_p_elem = OxmlElement("w:p")
        p_ordens._p.addnext(novo_p_elem)
        paragrafo_quebra = Paragraph(novo_p_elem, p_ordens._parent)
        paragrafo_quebra.add_run().add_break(WD_BREAK.PAGE)

    # O modelo original já vem com um parágrafo em branco + quebra de
    # página entre "3.4 Ordens de Serviço - Handover" e "4. CAPACITAÇÃO
    # DA EQUIPE" (espaço reservado pra colar as páginas da Fracttal à
    # mão, no fluxo manual antigo). Quando NÃO vamos mesclar o PDF da
    # Fracttal, essa quebra sobra e deixa uma página quase em branco no
    # meio do relatório — então removemos esse parágrafo nesse caso.
    if not inserir_quebra_para_fracttal:
        p_ordens = _achar_paragrafo(doc, "Ordens de Serviço - Handover")
        p_capacitacao = _achar_paragrafo(doc, "CAPACITAÇÃO DA EQUIPE")
        if p_ordens is not None and p_capacitacao is not None:
            atual = p_ordens._p.getnext()
            while atual is not None and atual is not p_capacitacao._p:
                proximo = atual.getnext()
                atual.getparent().remove(atual)
                atual = proximo

    # Punch List (seção nova em paisagem, no final)
    _adicionar_secao_punchlist(doc, dados.get("punchList", []))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Conversão pra PDF + merge com o PDF da Fracttal ─────────────────────

def _converter_docx_para_pdf(docx_bytes):
    """Converte bytes de um .docx pra PDF via LibreOffice headless. Cada
    chamada usa um perfil de usuário temporário isolado, pra não haver
    conflito entre requisições concorrentes na mesma VM."""
    with tempfile.TemporaryDirectory(prefix="handover_lo_") as tmpdir:
        docx_path = os.path.join(tmpdir, "entrada.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        perfil = os.path.join(tmpdir, "perfil_lo")
        os.makedirs(perfil, exist_ok=True)

        resultado = subprocess.run(
            ["soffice", "--headless", "--norestore",
             f"-env:UserInstallation=file://{perfil}",
             "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
            capture_output=True, timeout=90,
        )
        pdf_path = os.path.join(tmpdir, "entrada.pdf")
        if not os.path.exists(pdf_path):
            erro = (resultado.stderr or b"").decode(errors="ignore")[:800]
            raise RuntimeError(f"Falha ao converter o .docx para PDF (LibreOffice): {erro}")
        with open(pdf_path, "rb") as f:
            return f.read()


def _pagina_so_tem_cabecalho(texto):
    """Considera 'vazia' (pra fins do corte) uma página que só tem o
    cabeçalho padrão (logo/título/nº de página) e nenhum conteúdo de
    corpo — toda página tem o cabeçalho, então checar `if texto` sozinho
    NUNCA reconhece a página reservada como vazia (bug real, achado num
    teste com PDF da Fracttal anexado: a página reservada sobrava depois
    do PDF inserido porque o cabeçalho a fazia parecer 'com conteúdo')."""
    marcadores_cabecalho = ("relatório de handover", "operação", "de ativos",
                             "grid co.", "rev.00", "página")
    for linha in (texto or "").split("\n"):
        linha_normalizada = linha.strip().lower()
        if not linha_normalizada:
            continue
        if not any(m in linha_normalizada for m in marcadores_cabecalho):
            return False  # achou uma linha que não é do cabeçalho => tem conteúdo
    return True


def _dividir_pdf_no_marcador(pdf_bytes, texto_marcador="Ordens de Serviço - Handover"):
    """Divide o PDF em duas partes: tudo até a ÚLTIMA página que contém
    `texto_marcador` (inclusive) vira a parte 1, o resto vira a parte 2.
    Como gerar_handover_usina_docx (quando inserir_quebra_para_fracttal=True)
    insere uma quebra de página em branco logo depois desse título — só
    pra garantir que essa página não tenha conteúdo da seção 4 junto,
    protegendo o ponto de corte — essa página em branco cai sempre no
    INÍCIO da parte 2. Sem removê-la, ela sobra como página em branco
    DEPOIS do PDF da Fracttal ser inserido no meio (bug real, achado
    pelo Fred num teste com PDF anexado). Por isso: depois de cortar,
    remove páginas só-com-cabeçalho do início da parte 2."""
    reader = PdfReader(BytesIO(pdf_bytes))
    pagina_split = None
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            if texto_marcador in (page.extract_text() or ""):
                pagina_split = i

    indices_parte1, indices_parte2 = [], []
    for i in range(len(reader.pages)):
        (indices_parte1 if (pagina_split is None or i <= pagina_split) else indices_parte2).append(i)

    # remove páginas só-com-cabeçalho do INÍCIO da parte 2 (a quebra reservada)
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        while indices_parte2:
            texto = pdf.pages[indices_parte2[0]].extract_text() or ""
            if not _pagina_so_tem_cabecalho(texto):
                break
            indices_parte2.pop(0)

    parte1, parte2 = PdfWriter(), PdfWriter()
    for i in indices_parte1:
        parte1.add_page(reader.pages[i])
    for i in indices_parte2:
        parte2.add_page(reader.pages[i])

    buf1, buf2 = BytesIO(), BytesIO()
    parte1.write(buf1)
    parte2.write(buf2)
    buf1.seek(0)
    buf2.seek(0)
    return buf1.getvalue(), buf2.getvalue()


def gerar_handover_usina_completo(dados, fracttal_pdf_bytes=None):
    """
    Gera o Relatório de Handover completo.

    Sem PDF da Fracttal: retorna (bytes_docx, "docx") — o documento
    editável, direto do modelo real.

    Com PDF da Fracttal: converte o .docx pra PDF (LibreOffice), corta
    logo após o título "3.4 Ordens de Serviço - Handover", insere as
    páginas do PDF da Fracttal ali no meio, e devolve
    (bytes_pdf_final, "pdf") — documento + OS mesclados, prontos pra
    enviar ao cliente.
    """
    docx_bytes = gerar_handover_usina_docx(
        dados, inserir_quebra_para_fracttal=bool(fracttal_pdf_bytes)
    ).read()
    if not fracttal_pdf_bytes:
        return docx_bytes, "docx"

    pdf_bytes = _converter_docx_para_pdf(docx_bytes)
    parte1_bytes, parte2_bytes = _dividir_pdf_no_marcador(pdf_bytes)

    writer = PdfWriter()
    for chunk in (parte1_bytes, fracttal_pdf_bytes, parte2_bytes):
        if not chunk:
            continue
        for page in PdfReader(BytesIO(chunk)).pages:
            writer.add_page(page)

    saida = BytesIO()
    writer.write(saida)
    saida.seek(0)
    return saida.read(), "pdf"
