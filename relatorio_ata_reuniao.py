# -*- coding: utf-8 -*-
"""
relatorio_ata_reuniao.py
─────────────────────────────────────────────────────────────────────────
Geração de "Ata de Reunião" (.docx) no padrão visual Grid Co., a partir de
uma transcrição de reunião (Teams) estruturada pela IA (Gemini) em app.py.

Segue à risca o padrão descrito na skill "grid-relatorio-operacional"
(paleta, tipografia, layout de página, capa, section headers, step rows,
activity table e premissas) — mesma identidade visual usada nos relatórios
de referência da Grid Co., reimplementada aqui em python-docx (o backend
roda Flask/Python; o protótipo original desta skill usa docx.js/Node).
"""
import os
from io import BytesIO

from docx import Document
from docx.shared import Pt, Twips, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════ PALETA (exata da skill) ═══════════════════
GREEN = "A2CA40"
DARK_COVER = "1C1F3B"
GRAY_LIGHT = "F5F5F2"
GREEN_LIGHT = "EBF4D0"
GREEN_MED = "D6EFB3"
WARN_LIGHT = "FFF3CD"
WHITE = "FFFFFF"
GRAY_TEXT = "666666"
GREEN_RESP = "5E8C1A"
RESP_BODY = "555555"
NOTE_TEXT = "444444"
BORDER = "DDDDDD"
GRAY_LABEL = "888888"
FONT = "Arial"

ASSETS_DIR = os.path.join(os.path.dirname(__file__), "templates", "ata_reuniao_assets")
ICON_GREEN = os.path.join(ASSETS_DIR, "icon_green.png")

# largura de conteúdo padrão (DXA) — igual à skill
CONTENT_W = 9360


# ═══════════════════════════ HELPERS DE BAIXO NÍVEL (OOXML) ════════════

# Ordem exigida pelo schema OOXML (CT_TcPr) — elementos precisam ser
# inseridos nesta sequência dentro de <w:tcPr>, senão o Word considera o
# arquivo corrompido.
_TCPR_ORDER = ["cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders",
               "shd", "noWrap", "tcMar", "textDirection", "tcFitText",
               "vAlign", "hideMark"]

# Ordem exigida pelo schema OOXML (CT_PPr).
_PPR_ORDER = ["pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
              "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd",
              "tabs", "suppressAutoHyphens", "kinsoku", "wordWrap",
              "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
              "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
              "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc",
              "textDirection", "textAlignment", "textboxTightWrap",
              "outlineLvl", "divId", "cnfStyle", "rPr"]

# Ordem exigida pelo schema OOXML (CT_TblPr).
_TBLPR_ORDER = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
                "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
                "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
                "tblCellMar", "tblLook", "tblCaption", "tblDescription"]


def _ordered_insert(parent, new_el, order_list):
    """Insere new_el em parent respeitando a ordem do schema OOXML —
    procura o primeiro filho já existente cuja tag vem DEPOIS de new_el
    na sequência e insere logo antes dele; senão, insere no final."""
    new_tag = new_el.tag.split("}")[-1]
    try:
        new_idx = order_list.index(new_tag)
    except ValueError:
        parent.append(new_el)
        return
    for child in parent:
        child_tag = child.tag.split("}")[-1]
        if child_tag in order_list and order_list.index(child_tag) > new_idx:
            child.addprevious(new_el)
            return
    parent.append(new_el)


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _ordered_insert(tcPr, shd, _TCPR_ORDER)


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """CT_TcMar também usa 'start'/'end' em vez de 'left'/'right' neste
    schema (mesma convenção de CT_TcBorders)."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    _ordered_insert(tcPr, mar, _TCPR_ORDER)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Cada lado: None => sem borda (nil); dict(sz=, color=) => single line.
    CT_TcBorders (bordas de célula) usa "start"/"end" em vez de
    "left"/"right" no schema validado aqui (diferente de CT_PBdr, que usa
    left/right normalmente)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge_name, edge_data in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        tag = OxmlElement(f"w:{edge_name}")
        if edge_data is None:
            tag.set(qn("w:val"), "nil")
        else:
            tag.set(qn("w:val"), "single")
            tag.set(qn("w:sz"), str(edge_data.get("sz", 4)))
            tag.set(qn("w:color"), edge_data.get("color", BORDER))
        tcBorders.append(tag)
    _ordered_insert(tcPr, tcBorders, _TCPR_ORDER)


def _set_table_fixed_layout(table, widths_dxa):
    # table.autofit = False já insere <w:tblLayout w:type="fixed"/> na
    # posição correta do schema (python-docx cuida disso sozinho) — inserir
    # de novo manualmente aqui duplicava o elemento e quebrava o XML.
    table.autofit = False
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in widths_dxa:
        gridcol = OxmlElement("w:gridCol")
        gridcol.set(qn("w:w"), str(w))
        tblGrid.append(gridcol)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_dxa):
                cell.width = Twips(widths_dxa[idx])


def _set_row_height(row, dxa, rule="atLeast"):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(dxa))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def _set_paragraph_shading(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _ordered_insert(pPr, shd, _PPR_ORDER)


def _set_paragraph_left_border(paragraph, color=GREEN, sz=12, space=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pBdr.append(left)
    _ordered_insert(pPr, pBdr, _PPR_ORDER)


def _run(paragraph, text, color=DARK_COVER, bold=False, italic=False, size_half_pt=20, font=FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size_half_pt / 2)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _para(cell_or_doc, alignment=None, space_after=None, space_before=None, line=None):
    p = cell_or_doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after / 20)
    if space_before is not None:
        pf.space_before = Pt(space_before / 20)
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line / 240
    return p


def _new_table(doc_or_cell, rows, cols, widths_dxa):
    table = doc_or_cell.add_table(rows=rows, cols=cols)
    table.style = None
    _set_table_fixed_layout(table, widths_dxa)
    return table


# ═══════════════════════════ COMPONENTES DA SKILL ═══════════════════════

def _build_cover(doc, titulo_capa, subtitulo_capa, supertitulo, subtitulo_doc, data_extenso, rodape_doc):
    table = _new_table(doc, 1, 2, [339, 11578])
    row = table.rows[0]
    _set_row_height(row, 15220)

    green_cell, dark_cell = row.cells
    _set_cell_bg(green_cell, GREEN)
    _set_cell_border(green_cell)
    _set_cell_margins(green_cell)

    _set_cell_bg(dark_cell, DARK_COVER)
    _set_cell_border(dark_cell)
    _set_cell_margins(dark_cell, top=900, left=700, right=700, bottom=700)
    dark_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # limpa parágrafo padrão vazio das células
    green_cell.paragraphs[0].text = ""
    p0 = dark_cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    run_img = p0.add_run()
    if os.path.exists(ICON_GREEN):
        run_img.add_picture(ICON_GREEN, width=Inches(0.4), height=Inches(0.4))
    _run(p0, "  Grid Co.", color=WHITE, bold=True, size_half_pt=36)

    p1 = _para(dark_cell)
    _run(p1, "Operação & Manutenção", color=GREEN, size_half_pt=16)

    p2 = _para(dark_cell, space_before=2800)
    _run(p2, "")

    p3 = _para(dark_cell, space_after=80)
    _run(p3, "ATA DE REUNIÃO", color=GRAY_LABEL, size_half_pt=16)

    p4 = _para(dark_cell)
    _run(p4, titulo_capa, color=WHITE, size_half_pt=44)

    p5 = _para(dark_cell, space_after=40)
    _run(p5, subtitulo_capa, color=GREEN, bold=True, size_half_pt=72)

    p6 = _para(dark_cell)
    pPr = p6._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "20")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), GREEN)
    pBdr.append(bottom)
    _ordered_insert(pPr, pBdr, _PPR_ORDER)
    _run(p6, "")

    for _ in range(4):
        p = _para(dark_cell)
        _run(p, "", size_half_pt=15)

    p11 = _para(dark_cell, space_after=80)
    _run(p11, supertitulo, color=GRAY_LABEL, size_half_pt=15)

    p12 = _para(dark_cell)
    _run(p12, subtitulo_doc, color=WHITE, bold=True, size_half_pt=32)

    p13 = _para(dark_cell, space_before=40)
    _run(p13, data_extenso, color=GREEN, size_half_pt=26)

    p14 = _para(dark_cell)
    _run(p14, "", size_half_pt=15)

    p15 = _para(dark_cell)
    _run(p15, rodape_doc, color=GRAY_TEXT, size_half_pt=15)


def _sec_header(doc, texto):
    table = _new_table(doc, 1, 2, [160, 9200])
    c0, c1 = table.rows[0].cells
    _set_cell_bg(c0, GREEN)
    _set_cell_border(c0)
    _set_cell_margins(c0)
    c0.paragraphs[0].text = ""

    _set_cell_bg(c1, DARK_COVER)
    _set_cell_border(c1)
    _set_cell_margins(c1, top=120, bottom=120, left=200)
    p = c1.paragraphs[0]
    _run(p, texto, color=WHITE, bold=True, size_half_pt=24)
    return table


def _body_p(doc, texto, italic=False, space_after=120, size_half_pt=20):
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=space_after, line=270)
    _run(p, texto, italic=italic, size_half_pt=size_half_pt)
    return p


def _spacer(doc, after=120):
    p = _para(doc, space_after=after)
    _run(p, "")
    return p


_step_alt = {"v": False}


def _step_row(doc, n, titulo, responsavel, corpo, callout_label=None, callout_texto=None, callout_tipo="green"):
    fill = GRAY_LIGHT if _step_alt["v"] else WHITE
    _step_alt["v"] = not _step_alt["v"]

    table = _new_table(doc, 1, 2, [600, 8760])
    c0, c1 = table.rows[0].cells
    _set_cell_bg(c0, DARK_COVER)
    _set_cell_border(c0)
    _set_cell_margins(c0, top=120, bottom=120, left=100, right=100)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p0, str(n), color=WHITE, bold=True, size_half_pt=24)

    _set_cell_bg(c1, fill)
    _set_cell_border(c1, bottom={"sz": 2, "color": BORDER})
    _set_cell_margins(c1, top=120, bottom=120, left=200, right=160)

    p_titulo = c1.paragraphs[0]
    p_titulo.paragraph_format.space_after = Pt(2)
    _run(p_titulo, titulo, bold=True, size_half_pt=20)

    if responsavel:
        p_resp = _para(c1, space_after=60)
        _run(p_resp, "Responsável: ", color=GREEN_RESP, bold=True, size_half_pt=17)
        _run(p_resp, responsavel, color=RESP_BODY, size_half_pt=17)

    p_corpo = _para(c1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=60, line=270)
    _run(p_corpo, corpo, size_half_pt=19)

    if callout_label and callout_texto:
        p_call = _para(c1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=40, line=260)
        fill_callout = WARN_LIGHT if callout_tipo == "warn" else GREEN_MED
        _set_paragraph_shading(p_call, fill_callout)
        _run(p_call, callout_label + "  ", color=NOTE_TEXT, bold=True, size_half_pt=18)
        _run(p_call, callout_texto, color=NOTE_TEXT, size_half_pt=18)

    return table


def _activity_table(doc, itens):
    widths = [440, 5560, 3360]
    table = _new_table(doc, 1 + len(itens), 3, widths)

    hdr = table.rows[0].cells
    headers = ["#", "Ação / Encaminhamento", "Responsável / Prazo"]
    for i, cell in enumerate(hdr):
        _set_cell_bg(cell, DARK_COVER)
        _set_cell_border(cell)
        _set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        _run(p, headers[i], color=WHITE, bold=True, size_half_pt=18)

    for idx, item in enumerate(itens):
        fill = WHITE if idx % 2 == 0 else GRAY_LIGHT
        row_cells = table.rows[idx + 1].cells

        _set_cell_bg(row_cells[0], DARK_COVER)
        _set_cell_border(row_cells[0])
        _set_cell_margins(row_cells[0], top=80, bottom=80, left=100, right=100)
        _run(row_cells[0].paragraphs[0], str(idx + 1), color=WHITE, bold=True, size_half_pt=18)

        _set_cell_bg(row_cells[1], fill)
        _set_cell_border(row_cells[1])
        _set_cell_margins(row_cells[1], top=80, bottom=80, left=100, right=100)
        p1 = row_cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _run(p1, item.get("acao", ""), size_half_pt=18)

        _set_cell_bg(row_cells[2], fill)
        _set_cell_border(row_cells[2])
        _set_cell_margins(row_cells[2], top=80, bottom=80, left=100, right=100)
        _run(row_cells[2].paragraphs[0], item.get("resp", ""), color=GRAY_TEXT, italic=True, size_half_pt=18)

    return table


def _prem_row(doc, n, texto, alt):
    fill = GRAY_LIGHT if alt else WHITE
    table = _new_table(doc, 1, 2, [120, 9240])
    c0, c1 = table.rows[0].cells
    _set_cell_bg(c0, GREEN)
    _set_cell_border(c0)
    _set_cell_margins(c0)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p0, str(n), color=WHITE, bold=True, size_half_pt=18)

    _set_cell_bg(c1, fill)
    _set_cell_border(c1, bottom={"sz": 1, "color": BORDER})
    _set_cell_margins(c1, top=80, bottom=80, left=150, right=150)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p1, texto, size_half_pt=19)
    return table


def _data_table(doc, headers, linhas, widths_dxa):
    table = _new_table(doc, 1 + len(linhas), len(headers), widths_dxa)
    hdr = table.rows[0].cells
    for i, cell in enumerate(hdr):
        _set_cell_bg(cell, DARK_COVER)
        _set_cell_border(cell)
        _set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        _run(cell.paragraphs[0], headers[i], color=WHITE, bold=True, size_half_pt=18)

    for idx, linha in enumerate(linhas):
        fill = WHITE if idx % 2 == 0 else GRAY_LIGHT
        row_cells = table.rows[idx + 1].cells
        for i, val in enumerate(linha):
            _set_cell_bg(row_cells[i], fill)
            _set_cell_border(row_cells[i])
            _set_cell_margins(row_cells[i], top=80, bottom=80, left=100, right=100)
            _run(row_cells[i].paragraphs[0], str(val), size_half_pt=18)
    return table


def _add_header_footer(section, header_texto, footer_texto):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(3)
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BORDER)
    pBdr.append(bottom)
    _ordered_insert(pPr, pBdr, _PPR_ORDER)
    run_img = hp.add_run()
    if os.path.exists(ICON_GREEN):
        run_img.add_picture(ICON_GREEN, width=Inches(0.19), height=Inches(0.19))
    _run(hp, "   " + header_texto, color=GRAY_TEXT, size_half_pt=16)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(3)
    pPr2 = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), BORDER)
    pBdr2.append(top)
    _ordered_insert(pPr2, pBdr2, _PPR_ORDER)
    _run(fp, footer_texto, color=GRAY_TEXT, italic=True, size_half_pt=16)


# ═══════════════════════════ FUNÇÃO PRINCIPAL ═══════════════════════════

def gerar_ata_reuniao_docx(dados):
    """
    dados = {
      "titulo_capa": "Reunião Semanal",
      "subtitulo_capa": "RENOGRID & GRID CO.",
      "supertitulo": "VISÃO OPERACIONAL — ACOMPANHAMENTO SEMANAL",
      "subtitulo_doc": "Acompanhamento de Usinas Fotovoltaicas",
      "data_extenso": "23 de julho de 2026",
      "rodape_capa": "Documento de uso interno — Grid Co. / Renogrid",
      "header_titulo": "Reunião Semanal",
      "clientes_label": "Renogrid & Grid Co.",
      "footer_revisao": "Revisão 1  |  julho 2026",
      "objetivo": "texto...",
      "topicos": [
        {"titulo": "...", "responsavel": "...", "corpo": "...",
         "callout_label": "Ação:", "callout_texto": "...", "callout_tipo": "green"}
      ],
      "cronograma": {  # opcional — pode ser None
        "titulo": "CRONOGRAMA DE POD...",
        "descricao": "texto intro itálico",
        "headers": [...], "linhas": [[...], ...],
        "col_widths": [2600, 2400, 4360]  # opcional, soma=9360
      },
      "acoes": [{"acao": "...", "resp": "..."}],
      "premissas": ["...", "..."]
    }
    Retorna bytes do .docx pronto.
    """
    _step_alt["v"] = False
    doc = Document()

    # Corrige um detalhe do template padrão do python-docx: o elemento
    # <w:zoom> em settings.xml vem sem o atributo obrigatório w:percent,
    # o que reprova a validação de schema (Word/LibreOffice abrem normal,
    # mas o arquivo fica tecnicamente inválido).
    settings_el = doc.settings.element
    zoom_el = settings_el.find(qn("w:zoom"))
    if zoom_el is not None and zoom_el.get(qn("w:percent")) is None:
        zoom_el.set(qn("w:percent"), "100")

    # ── remove parágrafo em branco padrão do documento (será a 1a seção) ──
    sec0 = doc.sections[0]
    sec0.page_width = Twips(11906)
    sec0.page_height = Twips(16838)
    sec0.top_margin = Twips(0)
    sec0.bottom_margin = Twips(0)
    sec0.left_margin = Twips(0)
    sec0.right_margin = Twips(0)

    _build_cover(
        doc,
        dados.get("titulo_capa", "Reunião"),
        dados.get("subtitulo_capa", "GRID CO."),
        dados.get("supertitulo", "VISÃO OPERACIONAL — ACOMPANHAMENTO SEMANAL"),
        dados.get("subtitulo_doc", "Acompanhamento de Usinas Fotovoltaicas"),
        dados.get("data_extenso", ""),
        dados.get("rodape_capa", "Documento de uso interno — Grid Co."),
    )

    # ── nova seção pro conteúdo (margens normais + header/footer) ──
    sec1 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec1.page_width = Twips(11906)
    sec1.page_height = Twips(16838)
    sec1.top_margin = Twips(1080)
    sec1.bottom_margin = Twips(1080)
    sec1.left_margin = Twips(1080)
    sec1.right_margin = Twips(1080)

    header_titulo = dados.get("header_titulo", dados.get("titulo_capa", "Ata de Reunião"))
    clientes_label = dados.get("clientes_label", dados.get("subtitulo_capa", "Grid Co."))
    _add_header_footer(
        sec1,
        f"ATA DE REUNIÃO | {clientes_label} — {header_titulo}",
        f"Grid Co.  |  Ata de Reunião  |  {dados.get('footer_revisao', 'Revisão 1')}",
    )

    n_sec = 1

    # 1. OBJETIVO
    _sec_header(doc, f"{n_sec}.  OBJETIVO DO DOCUMENTO")
    _spacer(doc, 80)
    _body_p(doc, dados.get("objetivo", ""))
    _spacer(doc, 160)
    n_sec += 1

    # 2. TÓPICOS DISCUTIDOS
    _sec_header(doc, f"{n_sec}.  TÓPICOS DISCUTIDOS")
    _spacer(doc, 80)
    _body_p(doc, "Principais assuntos tratados durante a reunião, em ordem de discussão.",
            italic=True, space_after=160)
    for i, t in enumerate(dados.get("topicos", []), start=1):
        _step_row(
            doc, i, t.get("titulo", ""), t.get("responsavel", ""), t.get("corpo", ""),
            t.get("callout_label"), t.get("callout_texto"), t.get("callout_tipo", "green"),
        )
    _spacer(doc, 160)
    n_sec += 1

    # 3. CRONOGRAMA (opcional)
    cronograma = dados.get("cronograma")
    if cronograma and cronograma.get("linhas"):
        _sec_header(doc, f"{n_sec}.  {cronograma.get('titulo', 'CRONOGRAMA')}")
        _spacer(doc, 80)
        if cronograma.get("descricao"):
            _body_p(doc, cronograma["descricao"], italic=True, space_after=160)
        headers = cronograma.get("headers", [])
        widths = cronograma.get("col_widths") or [CONTENT_W // max(len(headers), 1)] * len(headers)
        _data_table(doc, headers, cronograma.get("linhas", []), widths)
        _spacer(doc, 160)
        n_sec += 1

    # N. AÇÕES E ENCAMINHAMENTOS
    _sec_header(doc, f"{n_sec}.  AÇÕES E ENCAMINHAMENTOS")
    _spacer(doc, 80)
    _activity_table(doc, dados.get("acoes", []))
    _spacer(doc, 160)
    n_sec += 1

    # N+1. PREMISSAS GERAIS
    _sec_header(doc, f"{n_sec}.  PREMISSAS E OBSERVAÇÕES GERAIS")
    _spacer(doc, 80)
    premissas = dados.get("premissas", [])
    for i, texto in enumerate(premissas):
        _prem_row(doc, i + 1, texto, alt=(i % 2 == 1))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
